import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import zipfile

# ----- Bahasa / Translation -----
LANGUAGES = {
    "id": {
        # ... semua entri bahasa Indonesia ...
    },
    "en": {
        # ... semua entri bahasa Inggris ...
    },
}

def t(key):
    lang = st.session_state.get("lang", "id")
    return LANGUAGES.get(lang, LANGUAGES["id"]).get(key, key)


# ----- Utils: Add Hyperlink & Styling -----
def add_hyperlink(paragraph, text, url):
    # ... kode seperti sebelumnya ...
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.set(qn("xml:space"), "preserve")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_paragraph_style(doc):
    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)


# ----- AUTH -----
def show_login():
    st.title(t("welcome"))
    st.markdown(t("login"))
    with st.form("login_form"):
        username = st.text_input(t("username"))
        password = st.text_input(t("password"), type="password")
        submitted = st.form_submit_button(t("login_button"))
        if submitted:
            if username == "admin" and password == "surat123":
                st.session_state["login_state"] = True
                st.session_state["username"] = username
                st.session_state["logout_message"] = False
                st.experimental_rerun()
            else:
                st.error(t("login_fail"))

def show_logout():
    if st.sidebar.button(t("logout_button")):
        st.session_state["logout_message"] = True
        st.session_state["login_state"] = False
        st.session_state["username"] = ""
        st.experimental_rerun()


# ----- DASHBOARD -----
def page_dashboard():
    st.title(t("dashboard_title"))
    st.markdown(f"{t('welcome')}, **{st.session_state['username']}**!")

    with st.expander(t("tips"), expanded=True):
        st.write(t("tips_content"))

    st.markdown("---")

    generate_log = st.session_state.get("generate_log", [])

    total_surat = len(generate_log)
    berhasil = sum(1 for item in generate_log if item["Status"].startswith("✅"))
    gagal = total_surat - berhasil

    template_tersedia = st.session_state.get("template_count", 1)
    data_peserta_terakhir = st.session_state.get("last_data_rows", 0)

    statistik_data = {
        "Statistik": [
            t("total_letters"),
            t("letters_success"),
            t("letters_failed"),
            t("templates_available"),
            t("last_data_rows"),
        ],
        "Jumlah": [
            total_surat,
            berhasil,
            gagal,
            template_tersedia,
            data_peserta_terakhir,
        ],
    }
    df_statistik = pd.DataFrame(statistik_data)
    st.table(df_statistik)

    st.markdown("---")

    st.markdown("### " + t("letters_success_vs_failed"))
    fig, ax = plt.subplots()
    ax.bar([t("letters_success"), t("letters_failed")], [berhasil, gagal], color=["green", "red"])
    ax.set_ylabel(t("total_letters"))
    ax.set_title(t("letters_success_vs_failed"))
    st.pyplot(fig)

    st.markdown("---")

    st.markdown("### " + t("percentage_letters"))
    fig2, ax2 = plt.subplots()
    if total_surat > 0:
        ax2.pie(
            [berhasil, gagal],
            labels=[t("letters_success"), t("letters_failed")],
            autopct="%1.1f%%",
            colors=["green", "red"],
            startangle=90,
            wedgeprops={"edgecolor": "black"},
        )
        ax2.axis("equal")
        st.pyplot(fig2)
    else:
        st.write(t("no_data"))

    st.markdown("---")

    st.markdown("### " + t("last_activity"))
    aktivitas = []
    for item in reversed(generate_log[-5:]):
        aktivitas.append({"Aktivitas": f"{t('generate_title')} untuk {item['Nama']}", "Status": item["Status"]})
    if aktivitas:
        df_aktivitas = pd.DataFrame(aktivitas)
        st.table(df_aktivitas)
    else:
        st.write(t("no_activity"))

    st.markdown("---")

    st.markdown(t("app_version"))
    st.markdown(t("no_maintenance"))


# ----- GENERATE SURAT -----
def generate_letters_with_progress(template_file, df, col_name, col_link):
    output_zip = BytesIO()
    log = []

    with zipfile.ZipFile(output_zip, "w") as zf:
        total = len(df)
        progress_bar = st.progress(0)
        status_text = st.empty()

        for idx, row in df.iterrows():
            try:
                tpl = DocxTemplate(template_file)
                tpl.render({"nama_penyelenggara": row[col_name], "short_link": "[short_link]"})
                temp_buf = BytesIO()
                tpl.save(temp_buf)
                temp_buf.seek(0)

                doc = Document(temp_buf)
                for p in doc.paragraphs:
                    if "[short_link]" in p.text:
                        parts = p.text.split("[short_link]")
                        p.clear()
                        if parts[0]:
                            run_before = p.add_run(parts[0])
                            run_before.font.name = "Arial"
                            run_before.font.size = Pt(12)
                        add_hyperlink(p, str(row[col_link]), str(row[col_link]))
                        if len(parts) > 1 and parts[1]:
                            run_after = p.add_run(parts[1])
                            run_after.font.name = "Arial"
                            run_after.font.size = Pt(12)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                set_paragraph_style(doc)

                final_buf = BytesIO()
                doc.save(final_buf)
                zf.writestr(f"{row[col_name]}.docx", final_buf.getvalue())
                log.append({"Nama": row[col_name], "Status": "✅ Berhasil"})
            except Exception as e:
                log.append({"Nama": row[col_name], "Status": f"❌ Gagal: {str(e)}"})

            progress = int((idx + 1) / total * 100)
            progress_bar.progress(progress)
            status_text.text(f"{t('processing_letters')} {idx + 1} / {total}")

    output_zip.seek(0)
    return output_zip, log

def page_generate():
    st.title(t("generate_title"))

    template_file = st.file_uploader(t("upload_template"), type="docx")
    data_file = st.file_uploader(t("upload_data"), type="xlsx")

    if template_file and data_file:
        df = pd.read_excel(data_file)
        st.success(f"{len(df)} baris data berhasil dimuat")
        st.dataframe(df)

        col_name = st.selectbox(t("select_name_col"), df.columns)
        col_link = st.selectbox(t("select_link_col"), df.columns)

        search_name = st.text_input(t("search_name"), "")
        filtered_names = df[df[col_name].astype(str).str.contains(search_name, case=False, na=False)][col_name].unique()
        selected_name = st.selectbox(t("select_name_preview"), filtered_names)

        st.session_state["df"] = df
        st.session_state["col_name"] = col_name
        st.session_state["col_link"] = col_link
        st.session_state["selected_name"] = selected_name
        st.session_state["template_file"] = template_file

        if st.session_state.get("show_preview", True) and selected_name:
            row = df[df[col_name] == selected_name].iloc[0]
            tpl = DocxTemplate(template_file)
            tpl.render({"nama_penyelenggara": row[col_name], "short_link": "[short_link]"})
            temp_buf = BytesIO()
            tpl.save(temp_buf)
            temp_buf.seek(0)

            doc = Document(temp_buf)
            preview_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            st.text_area(t("preview_letter"), preview_text, height=300)

            preview_buf = BytesIO()
            doc.save(preview_buf)
            preview_buf.seek(0)

            st.download_button(
                label=f"{t('download_preview')} ({row[col_name]})",
                data=preview_buf.getvalue(),
                file_name=f"preview_{row[col_name]}.docx",
            )

        if st.button(t("generate_all")):
            with st.spinner(t("processing_letters")):
                zip_file, log = generate_letters_with_progress(template_file, df, col_name, col_link)
            st.success(t("generate_done"))
            st.download_button(t("download_all_zip"), zip_file.getvalue(), file_name="surat_massal.zip")
            with st.expander(t("view_log")):
                st.dataframe(pd.DataFrame(log))
    else:
        st.info(t("upload_first"))


# ----- ANALISA DATA -----
def page_analysis():
    st.title(t("analysis_title"))

    data_file = st.file_uploader(t("upload_data"), type=["xlsx", "csv"])
    if data_file:
        try:
            if data_file.name.endswith(".csv"):
                df = pd.read_csv(data_file)
            else:
                df = pd.read_excel(data_file)
        except Exception as e:
            st.error(f"Error membaca file: {e}")
            return

        st.success(f"Data berhasil dimuat, {len(df)} baris.")
        st.dataframe(df)

        st.markdown("### Statistik Deskriptif")
        st.write(df.describe())

        numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
        if numeric_cols:
            col = st.selectbox("Pilih kolom untuk histogram", numeric_cols)
            bins = st.slider("Jumlah bin", 5, 100, 20)

            fig, ax = plt.subplots()
            ax.hist(df[col].dropna(), bins=bins, color="skyblue", edgecolor="black")
            ax.set_title(f"Histogram kolom {col}")
            ax.set_xlabel(col)
            ax.set_ylabel("Frekuensi")
            st.pyplot(fig)
        else:
            st.info("Tidak ada kolom numerik untuk histogram.")
    else:
        st.info(t("upload_first"))


# ----- MAIN APP -----
def check_session_timeout():
    from datetime import datetime, timedelta
    SESSION_TIMEOUT = timedelta(minutes=15)
    if "last_active" in st.session_state:
        if datetime.now() - st.session_state.last_active > SESSION_TIMEOUT:
            st.session_state.clear()
            st.experimental_rerun()
    st.session_state["last_active"] = datetime.now()


def show_main_app():
    check_session_timeout()
    show_logout()

    st.sidebar.title(t("choose_language"))
    lang = st.sidebar.selectbox(
        "",
        ["id", "en"],
        index=0 if st.session_state.get("lang", "id") == "id" else 1,
        format_func=lambda x: "Indonesia" if x == "id" else "English",
    )
    st.session_state["lang"] = lang

    st.sidebar.title("Menu")
    page = st.sidebar.radio(
        "Navigasi",
        [
            t("dashboard_title"),
            t("generate_title"),
            t("analysis_title"),
        ],
    )

    if page == t("dashboard_title"):
        page_dashboard()
    elif page == t("generate_title"):
        page_generate()
    else:
        page_analysis()


if __name__ == "__main__":
    if "login_state" not in st.session_state:
        st.session_state["login_state"] = False
    if "lang" not in st.session_state:
        st.session_state["lang"] = "id"

    if st.session_state.get("logout_message", False):
        st.title(t("logout_msg"))
        st.markdown(t("logout_submsg"))
        if st.button(t("back_login")):
            st.session_state["logout_message"] = False
            st.experimental_rerun()
    elif st.session_state["login_state"]:
        show_main_app()
    else:
        show_login()
