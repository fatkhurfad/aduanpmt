import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from io import BytesIO
import zipfile
from modules.utils import add_hyperlink, set_paragraph_style
from modules.config import t
from docx.shared import Pt

def generate_letters_with_progress(template_file, df, col_name, col_link):
    output_zip = BytesIO()
    log = []
    with zipfile.ZipFile(output_zip, "w") as zf:
        total = len(df)
        progress_bar = st.progress(0)
        status_text = st.empty()
        for idx, row in df.iterrows():
            try:
                tpl = DocxTemplate(template_file)
                tpl.render({"nama_penyelenggara": row[col_name], "short_link": "[short_link]"})
                temp_buf = BytesIO()
                tpl.save(temp_buf)
                temp_buf.seek(0)

                doc = Document(temp_buf)
                for p in doc.paragraphs:
                    if "[short_link]" in p.text:
                        parts = p.text.split("[short_link]")
                        p.clear()
                        if parts[0]:
                            run_before = p.add_run(parts[0])
                            run_before.font.name = "Arial"
                            run_before.font.size = Pt(12)
                        add_hyperlink(p, str(row[col_link]), str(row[col_link]))
                        if len(parts) > 1 and parts[1]:
                            run_after = p.add_run(parts[1])
                            run_after.font.name = "Arial"
                            run_after.font.size = Pt(12)
                    p.alignment = 1  # WD_ALIGN_PARAGRAPH.JUSTIFY

                set_paragraph_style(doc)

                final_buf = BytesIO()
                doc.save(final_buf)
                zf.writestr(f"{row[col_name]}.docx", final_buf.getvalue())
                log.append({"Nama": row[col_name], "Status": "✅ Berhasil"})
            except Exception as e:
                log.append({"Nama": row[col_name], "Status": f"❌ Gagal: {str(e)}"})

            progress = int((idx + 1) / total * 100)
            progress_bar.progress(progress)
            status_text.text(f"{t('processing_letters', st.session_state.lang)} {idx + 1} / {total}")

    output_zip.seek(0)
    return output_zip, log

def page_generate():
    st.title(t("generate_title", st.session_state.lang))

    template_file = st.file_uploader(t("upload_template", st.session_state.lang), type="docx")
    data_file = st.file_uploader(t("upload_data", st.session_state.lang), type="xlsx")

    if template_file and data_file:
        df = pd.read_excel(data_file)
        st.success(f"{len(df)} rows loaded successfully")
        st.dataframe(df)

        col_name = st.selectbox(t("select_name_col", st.session_state.lang), df.columns)
        col_link = st.selectbox(t("select_link_col", st.session_state.lang), df.columns)

        search_name = st.text_input(t("search_name", st.session_state.lang), "")
        filtered_names = df[df[col_name].astype(str).str.contains(search_name, case=False, na=False)][col_name].unique()
        selected_name = st.selectbox(t("select_name_preview", st.session_state.lang), filtered_names)

        st.session_state.df = df
        st.session_state.col_name = col_name
        st.session_state.col_link = col_link
        st.session_state.selected_name = selected_name
        st.session_state.template_file = template_file

        if st.session_state.get("show_preview", True) and selected_name:
            row = df[df[col_name] == selected_name].iloc[0]
            tpl = DocxTemplate(template_file)
            tpl.render({"nama_penyelenggara": row[col_name], "short_link": "[short_link]"})
            temp_buf = BytesIO()
            tpl.save(temp_buf)
            temp_buf.seek(0)

            doc = Document(temp_buf)
            preview_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            st.text_area(t("preview_letter", st.session_state.lang), preview_text, height=300)

            preview_buf = BytesIO()
            doc.save(preview_buf)
            preview_buf.seek(0)

            st.download_button(
                label=f"{t('download_preview', st.session_state.lang)} ({row[col_name]})",
                data=preview_buf.getvalue(),
                file_name=f"preview_{row[col_name]}.docx",
            )

        if st.button(t("generate_all", st.session_state.lang)):
            with st.spinner(t("processing_letters", st.session_state.lang)):
                zip_file, log = generate_letters_with_progress(template_file, df, col_name, col_link)
            st.success(t("generate_done", st.session_state.lang))
            st.download_button(t("download_all_zip", st.session_state.lang), zip_file.getvalue(), file_name="surat_massal.zip")
            with st.expander(t("view_log", st.session_state.lang)):
                st.dataframe(pd.DataFrame(log))
    else:
        st.info(t("upload_first", st.session_state.lang))